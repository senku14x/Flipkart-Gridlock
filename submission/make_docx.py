"""
submission/make_docx.py: render the markdown submission docs into .docx (Word).

  - ParkPulse_Pitch_Deck.docx  : one slide per page (title, bullets, speaker notes,
                                 visual cue), built from PITCH_DECK.md.
  - ParkPulse_Solution.docx    : the solution doc with headings, bullets, and tables,
                                 built from SOLUTION.md.

Run:  pip install python-docx && python submission/make_docx.py
"""
from __future__ import annotations
import os
import re

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = os.path.dirname(os.path.abspath(__file__))
ACCENT = RGBColor(0xE4, 0x57, 0x2E)
INK = RGBColor(0x1B, 0x2A, 0x4A)
GREY = RGBColor(0x55, 0x5F, 0x70)

INLINE = re.compile(r"(\*\*.+?\*\*|`.+?`|\*[^*].*?\*)")


def add_runs(p, text):
    """Add a paragraph's text, honouring **bold**, `code`, *italic*."""
    pos = 0
    for m in INLINE.finditer(text):
        if m.start() > pos:
            p.add_run(text[pos:m.start()])
        tok = m.group(0)
        if tok.startswith("**"):
            p.add_run(tok[2:-2]).bold = True
        elif tok.startswith("`"):
            r = p.add_run(tok[1:-1]); r.font.name = "Consolas"; r.font.color.rgb = INK
        else:
            p.add_run(tok[1:-1]).italic = True
        pos = m.end()
    if pos < len(text):
        p.add_run(text[pos:])


def base_style(doc):
    n = doc.styles["Normal"]
    n.font.name = "Calibri"
    n.font.size = Pt(11)


# ----------------------------------------------------------------------------- deck
def build_deck(src, out):
    doc = Document(); base_style(doc)
    text = open(src, encoding="utf-8").read()
    # everything before the first slide is the cover note
    head, _, body = text.partition("\n### ")
    body = "### " + body
    title = head.splitlines()[0].lstrip("# ").strip()

    # cover page
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("ParkPulse"); r.bold = True; r.font.size = Pt(44); r.font.color.rgb = ACCENT
    p2 = doc.add_paragraph()
    r2 = p2.add_run("Find the parking that actually chokes traffic, and patrol it first.")
    r2.font.size = Pt(16); r2.font.color.rgb = INK
    p3 = doc.add_paragraph()
    r3 = p3.add_run("Gridlock Hackathon 2.0 (Flipkart), Round 2  ·  pitch deck")
    r3.font.size = Pt(12); r3.font.color.rgb = GREY
    doc.add_page_break()

    slides = re.split(r"\n(?=### )", body)
    for s in slides:
        lines = s.strip().splitlines()
        if not lines or not lines[0].startswith("### "):
            continue
        heading = lines[0][4:].strip()
        num, _, name = heading.partition(". ")
        # slide label + title
        lab = doc.add_paragraph(); rl = lab.add_run(f"SLIDE {num}")
        rl.font.size = Pt(10); rl.bold = True; rl.font.color.rgb = ACCENT
        tt = doc.add_paragraph(); rt = tt.add_run(name or heading)
        rt.bold = True; rt.font.size = Pt(26); rt.font.color.rgb = INK
        doc.add_paragraph()

        for ln in lines[1:]:
            ln = ln.rstrip()
            if not ln or ln == "---":
                continue
            if ln.startswith("- "):
                p = doc.add_paragraph(style="List Bullet"); add_runs(p, ln[2:])
            elif ln.startswith("> Speaker:"):
                lp = doc.add_paragraph(); lr = lp.add_run("Speaker notes")
                lr.bold = True; lr.font.size = Pt(10); lr.font.color.rgb = GREY
                sp = doc.add_paragraph(); sr = sp.add_run(ln[len("> Speaker:"):].strip())
                sr.italic = True; sr.font.color.rgb = GREY
            elif ln.startswith("> "):
                q = doc.add_paragraph(); qr = q.add_run(ln[2:].strip())
                qr.italic = True; qr.font.color.rgb = GREY
            elif ln.startswith("*Visual:*"):
                v = doc.add_paragraph()
                vr = v.add_run("Visual: "); vr.bold = True; vr.font.size = Pt(10); vr.font.color.rgb = GREY
                add_runs(v, ln[len("*Visual:*"):].strip())
                for r in v.runs[1:]:
                    r.font.size = Pt(10); r.font.color.rgb = GREY; r.italic = True
            else:
                add_runs(doc.add_paragraph(), ln)
        doc.add_page_break()
    doc.save(out)
    return len(slides)


# ------------------------------------------------------------------------- solution
def flush_table(doc, rows):
    if not rows:
        return
    header = [c.strip() for c in rows[0].strip("|").split("|")]
    body = [r for r in rows[2:]]  # skip the |---| separator
    t = doc.add_table(rows=1, cols=len(header)); t.style = "Light Grid Accent 1"
    for i, h in enumerate(header):
        cell = t.rows[0].cells[i]; cell.paragraphs[0].clear()
        add_runs(cell.paragraphs[0], h)
        for r in cell.paragraphs[0].runs:
            r.bold = True
    for row in body:
        cells = [c.strip() for c in row.strip("|").split("|")]
        tr = t.add_row().cells
        for i, c in enumerate(cells[:len(header)]):
            tr[i].paragraphs[0].clear(); add_runs(tr[i].paragraphs[0], c)
    doc.add_paragraph()


def build_solution(src, out):
    doc = Document(); base_style(doc)
    tbl = []
    for raw in open(src, encoding="utf-8"):
        ln = raw.rstrip("\n")
        if ln.startswith("|"):
            tbl.append(ln); continue
        if tbl:
            flush_table(doc, tbl); tbl = []
        if not ln.strip() or ln.strip() == "---":
            continue
        if ln.startswith("# "):
            h = doc.add_heading(level=0); r = h.add_run(ln[2:].strip()); r.font.color.rgb = ACCENT
        elif ln.startswith("## "):
            doc.add_heading(re.sub(r"[`*]", "", ln[3:].strip()), level=1)
        elif ln.startswith("### "):
            doc.add_heading(re.sub(r"[`*]", "", ln[4:].strip()), level=2)
        elif ln.startswith("- "):
            add_runs(doc.add_paragraph(style="List Bullet"), ln[2:])
        elif re.match(r"\d+\. ", ln):
            add_runs(doc.add_paragraph(style="List Number"), ln.split(". ", 1)[1])
        elif ln.startswith("> "):
            p = doc.add_paragraph(); p.style = "Intense Quote"; add_runs(p, ln[2:])
        else:
            add_runs(doc.add_paragraph(), ln)
    if tbl:
        flush_table(doc, tbl)
    doc.save(out)


def main():
    n = build_deck(os.path.join(HERE, "PITCH_DECK.md"),
                   os.path.join(HERE, "ParkPulse_Pitch_Deck.docx"))
    build_solution(os.path.join(HERE, "SOLUTION.md"),
                   os.path.join(HERE, "ParkPulse_Solution.docx"))
    build_solution(os.path.join(HERE, "VIDEO_SCRIPT.md"),
                   os.path.join(HERE, "ParkPulse_Video_Script.docx"))
    print(f"wrote ParkPulse_Pitch_Deck.docx ({n} slides) + ParkPulse_Solution.docx "
          "+ ParkPulse_Video_Script.docx")


if __name__ == "__main__":
    main()
